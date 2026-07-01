from pathlib import Path
from docx import Document


def clear_body(doc: Document):
    body = doc._body._element
    for child in list(body):
        body.remove(child)


def add_paragraph(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def split_report_sections(text: str):
    major_news = ""
    market_summary = ""

    if "二、市場摘要" in text:
        before_market, market_part = text.split("二、市場摘要", 1)
        market_summary = market_part.strip()
    else:
        before_market = text

    if "一、重大新聞" in before_market:
        _, major_part = before_market.split("一、重大新聞", 1)
        major_news = major_part.strip()
    else:
        major_news = before_market.strip()

    return major_news, market_summary


def add_content_block(doc, content: str):
    lines = [line.strip() for line in content.splitlines()]

    for line in lines:
        if not line:
            continue

        if line.startswith("【") and line.endswith("】"):
            add_paragraph(doc, line, style="Normal")

        elif line.startswith("●"):
            add_paragraph(doc, line.replace("●", "").strip(), style="List Paragraph")

        else:
            add_paragraph(doc, line, style="Normal")


def export_morning_report(
    txt_path: str,
    template_path: str,
    output_path: str,
    report_date: str
):
    txt_path = Path(txt_path)
    template_path = Path(template_path)
    output_path = Path(output_path)

    if not txt_path.exists():
        raise FileNotFoundError(f"找不到文字檔：{txt_path}")

    if not template_path.exists():
        raise FileNotFoundError(f"找不到 Word 範本：{template_path}")

    text = txt_path.read_text(encoding="utf-8")

    major_news, market_summary = split_report_sections(text)

    doc = Document(template_path)

    clear_body(doc)

    add_paragraph(doc, "早報資料", style="Normal")
    add_paragraph(doc, f"調撥科洪婉琪{report_date}", style="Normal")

    add_paragraph(doc, "重大新聞", style="List Paragraph")
    add_content_block(doc, major_news)

    add_paragraph(doc, "二、市場摘要", style="Normal")
    add_content_block(doc, market_summary)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return str(output_path)
