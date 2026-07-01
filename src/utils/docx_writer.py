from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


FONT_NAME = "標楷體"


def _set_font(run, size=14, bold=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def _new_document(font_size=14):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(font_size)

    return doc


# =========================================
# 會議紀錄
# =========================================

def write_signoff_docx(text: str, output_path: str):

    doc = _new_document(font_size=14)

    for line in text.splitlines():

        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)

        _set_font(run, size=14)

    doc.save(output_path)


# =========================================
# 晨報
# =========================================

def write_morning_docx(text: str, output_path: str):

    doc = _new_document(font_size=12)

    for line in text.splitlines():

        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()

        # 第一層標題
        if line.startswith("一、") or line.startswith("二、"):

            run = p.add_run(line)
            _set_font(run, size=12, bold=True)

        # 第二層標題
        elif line.startswith("●"):

            run = p.add_run(line)
            _set_font(run, size=12, bold=True)

        # 新聞標題
        elif line.startswith("【") and line.endswith("】"):

            run = p.add_run(line)
            _set_font(run, size=12, bold=True)

        # 內文
        else:

            run = p.add_run(line)
            _set_font(run, size=12)

    doc.save(output_path)
